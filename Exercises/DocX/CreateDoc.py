import docx
import os
os.chdir('c:\\users\\alan\\documents')
d = doc.Document()
d.paragraphs
d.paragraphs[0].text
p = d.paragraphs[1].text
p.runs()
p.runs[0].text
p.runs[1].text
p.runs[2].text
p.runs[3].text
p.runs[1].bold == True
p.runs[3].italic == True
p.runs[3].underline = True
p.runs[3].text = 'italic and underline.'
d.save('demo2.docx')
p.style = 'Title'
d.save('demo3.docx')