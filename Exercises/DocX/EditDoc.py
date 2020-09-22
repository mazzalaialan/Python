import docx
import os
os.chdir('c:\\users\\alan\\documents')
d = docx.Document('demo.docx')
d.add_paragraph('Hola, este es el primer parrafo')
d.add_paragraph('Este es el segundo parrafo')
d.save('demo4.docx')
p = d.paragraphs[0]
p.add_run('este es un nuevo run')
p.runs()
p.runs[1].bold = True
d.save('demo5 .docx')